# -*- coding: utf-8 -*-
"""
Genera, a partir de los resultados del modelo (OUTPUT/CASES_DATA y OUTPUT/LOCATION),
un documento Word con todas las tablas y figuras del capitulo de resultados del UC.
Portable: localiza las carpetas desde su propia ubicacion (UC MODEL/GRAFICAS Y TABLAS).
"""
import os, numpy as np, pandas as pd
import matplotlib; matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))  # .../UC MODEL
D   = os.path.join(BASE, "OUTPUT", "CASES_DATA")
LOC = os.path.join(BASE, "OUTPUT", "LOCATION")
IMG = os.path.join(os.path.dirname(os.path.abspath(__file__)), "figuras")
os.makedirs(IMG, exist_ok=True)
plt.rcParams.update({"font.size": 9, "axes.grid": True, "grid.alpha": 0.3, "figure.dpi": 130, "savefig.bbox": "tight"})
AZUL = RGBColor(0x1F, 0x3B, 0x57)

gens = pd.read_csv(os.path.join(D, "gens.csv")); NG = len(gens)

doc = Document()
doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(11)

def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexc); tcPr.append(sh)

def H(txt, lvl=1):
    h = doc.add_heading(txt, level=lvl)
    for r in h.runs: r.font.color.rgb = AZUL
    return h

def cap(txt):
    p = doc.add_paragraph(); r = p.add_run(txt); r.italic = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x55,0x55,0x55); p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_table(headers, rows, fontsz=8.0, hl=None):
    hl = hl or set()
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ""; pr = c.paragraphs[0]; run = pr.add_run(str(h))
        run.bold = True; run.font.size = Pt(fontsz); run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER; shade(c, "1F3B57")
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""; pr = cells[j].paragraphs[0]; run = pr.add_run(str(val)); run.font.size = Pt(fontsz)
            pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i in hl: shade(cells[j], "D9E8F5"); run.bold = True
            elif i % 2 == 1: shade(cells[j], "F2F6FA")
    return t

def img(path, w=6.2):
    doc.add_picture(path, width=Inches(w)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def fmt(v, d=2):
    """Formato español, sin miles ni ceros decimales no significativos.

    ``d`` representa el máximo de decimales visibles. Por ejemplo:
    0.0 -> 0; 30.0 -> 30; 4.4 -> 4,4; 99.01 -> 99,01.
    """
    try:
        value = round(float(v), d)
        if value == 0:
            return "0"
        text = f"{value:.{d}f}"
        if d > 0:
            text = text.rstrip("0").rstrip(".")
        return text.replace(".", ",")
    except: return str(v)

def load(case):
    return (pd.read_csv(os.path.join(D, f"{case}_gen.csv")),
            pd.read_csv(os.path.join(D, f"{case}_flujomax.csv")),
            pd.read_csv(os.path.join(D, f"{case}_obj.csv")),
            pd.read_csv(os.path.join(D, f"{case}_dim.csv")))

# ---------- generadores de TABLAS ----------
def tabla_unidades(g, use_wind, use_bess):
    headers = ["Hora"] + [f"G{i+1}" for i in range(NG)]
    if use_bess: headers += ["BESS neto [MW]"]
    if use_wind: headers += ["Eólico utilizado [MW]"]
    headers += ["Capacidad Térmica Comprometida [MW]"]
    rows = []
    for _, r in g.iterrows():
        line = [int(r.Hora)] + ["X" if int(r[f"u{i+1}"]) == 1 else "" for i in range(NG)]
        # Cap_comprometida proviene exclusivamente de las unidades térmicas:
        # sum(Pmax[g] * u[g,t]). El BESS y la eólica se reportan aparte y no se suman.
        if use_bess: line.append(fmt(r.BESS_net, 2))
        if use_wind: line.append(fmt(r.Eolica_uso, 2))
        line.append(fmt(r.Cap_comprometida, 2))
        rows.append(line)
    add_table(headers, rows, fontsz=7.0)

def tabla_flujomax(fm):
    add_table(["Hora","Línea más cargada","Flujo (MW)","Límite (MW)","Carga (%)"],
              [[int(r.Hora), r.Linea, fmt(r.Flujo,2), fmt(r.Limite,2), fmt(r.Carga_pct,2)] for _,r in fm.iterrows()], fontsz=8)

def tabla_obj(ob):
    desc = {"Cgen":"Costo variable de generación térmica","Cfijo":"Costo fijo de unidades comprometidas",
            "Carr":"Costo de arranque de unidades","Cvert":"Penalización por vertimiento eólico",
            "CSOC":"Penalización por SOC bajo del BESS","CBESS":"Operación y degradación del BESS",
            "Total":"Costo total de operación"}
    total = float(ob[ob.Componente == "Total"].Valor.iloc[0]); rows = []; hl = set()
    for i, r in ob.iterrows():
        if not bool(r.Aplica): rows.append([r.Componente, desc[r.Componente], "N/A", "—"])
        else:
            pct = 100*float(r.Valor)/total if total else 0
            rows.append([r.Componente, desc[r.Componente], fmt(r.Valor), fmt(pct,2)])
        if r.Componente == "Total": hl.add(i)
    add_table(["Componente","Descripción","Valor (USD)","% del total"], rows, fontsz=8.5, hl=hl)

def tabla_dim(dm):
    d = dm.set_index("Metrica").Valor
    rows = [["Variables continuas", fmt(int(d["Var_continuas"]),0)],
            ["Variables binarias", fmt(int(d["Var_binarias"]),0)],
            ["Variables totales", fmt(int(d["Var_totales"]),0)],
            ["Restricciones", fmt(int(d["Restricciones"]),0)],
            ["Estado de la solución", d["Estado"]],
            ["Brecha de optimalidad (gap)", f"{fmt(float(d['Gap_pct']),4)} %"],
            ["Valor objetivo óptimo (USD)", fmt(float(d["Costo_total"]))]]
    add_table(["Indicador","Valor"], rows, fontsz=9.5)

def tabla_eolica(g):
    add_table(["Hora","Disponible (MW)","Utilizada (MW)","Vertimiento (MW)"],
              [[int(r.Hora), fmt(r.Eolica_disp,2), fmt(r.Eolica_uso,2), fmt(r.Vertimiento,2)] for _,r in g.iterrows()], fontsz=8)

def tabla_bess(g):
    add_table(["Hora","Carga (MW)","Descarga (MW)","SOC (%)"],
              [[int(r.Hora), fmt(r.Charge,2), fmt(r.Discharge,2), fmt(r.SOC_pct,2)] for _,r in g.iterrows()], fontsz=8)

def tabla_cambios_compromiso(g_base, g_escenario, nombre_escenario):
    """Compara u[g,t] contra el escenario tradicional y agrupa horas consecutivas."""
    registros = []
    for i in range(NG):
        col = f"u{i+1}"
        base = g_base[col].round().astype(int).to_numpy()
        escenario = g_escenario[col].round().astype(int).to_numpy()
        horas = g_escenario.Hora.astype(int).to_numpy()
        indices = np.flatnonzero(base != escenario)
        if len(indices) == 0:
            continue

        inicio = indices[0]
        anterior = indices[0]
        estado_base = base[inicio]
        estado_escenario = escenario[inicio]

        for idx in list(indices[1:]) + [None]:
            consecutivo = (idx is not None and idx == anterior + 1)
            mismos_estados = (
                idx is not None
                and base[idx] == estado_base
                and escenario[idx] == estado_escenario
            )
            if consecutivo and mismos_estados:
                anterior = idx
                continue

            hora_ini, hora_fin = int(horas[inicio]), int(horas[anterior])
            periodo = str(hora_ini) if hora_ini == hora_fin else f"{hora_ini}–{hora_fin}"
            cambio = "Encendido" if estado_base == 0 and estado_escenario == 1 else "Apagado"
            registros.append([
                hora_ini,
                i + 1,
                periodo,
                f"G{i+1}",
                int(gens.iloc[i].Bus),
                int(estado_base),
                int(estado_escenario),
                cambio,
            ])

            if idx is not None:
                inicio = anterior = idx
                estado_base = base[idx]
                estado_escenario = escenario[idx]

    registros.sort(key=lambda x: (x[0], x[1]))
    if registros:
        rows = [r[2:] for r in registros]
    else:
        rows = [["—", "—", "—", "—", "—", "Sin cambios"]]

    add_table(
        ["Horas", "Generador", "Barra", "Estado tradicional",
         f"Estado {nombre_escenario}", "Cambio observado"],
        rows,
        fontsz=7.5,
    )

# ---------- generadores de FIGURAS ----------
def fig_despacho(case, g, titulo, use_wind=False, use_bess=False):
    """Grafica el despacho horario de todos los recursos del escenario.

    El apilamiento positivo contiene G1...G12, eólica utilizada y descarga del
    BESS. La carga del BESS se representa como una barra negativa. Esto no
    altera la columna Capacidad Térmica Comprometida de las tablas, que sigue
    calculándose exclusivamente como sum(Pmax[g] * u[g,t]).
    """
    horas = g.Hora.values; fig, ax = plt.subplots(figsize=(7.6,3.6)); bottom = np.zeros(len(horas)); cmap = plt.get_cmap("tab20")
    for i in range(NG):
        vals = g[f"p{i+1}"].values
        if np.allclose(vals,0): continue
        ax.bar(horas, vals, bottom=bottom, color=cmap(i%20), label=f"G{i+1}", width=0.85, linewidth=0); bottom += vals

    if use_wind:
        vals_eol = g.Eolica_uso.values
        ax.bar(horas, vals_eol, bottom=bottom, color="#2ECC71", label="Eólica utilizada",
               width=0.85, linewidth=0, edgecolor="#1E8449")
        bottom += vals_eol

    if use_bess:
        vals_desc = g.Discharge.values
        ax.bar(horas, vals_desc, bottom=bottom, color="#E67E22", label="BESS descarga",
               width=0.85, linewidth=0, edgecolor="#A04000")
        bottom += vals_desc
        ax.bar(horas, -g.Charge.values, color="#2980B9", label="BESS carga (negativa)",
               width=0.85, linewidth=0, hatch="//", alpha=0.85)

    ax.plot(horas, g.Demanda.values, color="#D62728", linestyle="-", marker="o",
            ms=4, lw=2.4, markeredgecolor="white", markeredgewidth=0.6,
            label="Demanda del sistema", zorder=10)
    if use_bess:
        ax.plot(horas, g.Demanda.values + g.Charge.values, color="#17202A",
                linestyle="--", lw=1.5, label="Demanda + carga BESS", zorder=10)
    ax.axhline(0, color="#555555", linewidth=0.7)
    ax.set_xlabel("Hora"); ax.set_ylabel("Potencia (MW)"); ax.set_title(titulo); ax.set_xticks(range(1,25))
    ax.legend(ncol=5, fontsize=6.5, loc="upper center", bbox_to_anchor=(0.5,-0.18))
    fn = os.path.join(IMG, f"{case}_despacho.png"); plt.savefig(fn); plt.close(); return fn

def fig_eolica(case, g):
    fig, ax = plt.subplots(figsize=(7.4,3.2))
    ax.plot(g.Hora, g.Eolica_disp, "-o", ms=3, color="#27AE60", label="Disponible")
    ax.bar(g.Hora, g.Eolica_uso, color="#82E0AA", label="Utilizada", width=0.7)
    ax.bar(g.Hora, g.Vertimiento, bottom=g.Eolica_uso, color="#E74C3C", alpha=0.7, label="Vertimiento", width=0.7)
    ax.set_xlabel("Hora"); ax.set_ylabel("Potencia (MW)"); ax.set_title("Energía eólica: disponible, utilizada y vertimiento")
    ax.set_xticks(range(1,25)); ax.legend(fontsize=8)
    fn = os.path.join(IMG, f"{case}_eolica.png"); plt.savefig(fn); plt.close(); return fn

def fig_bess(case, g):
    fig, ax1 = plt.subplots(figsize=(7.4,3.2))
    ax1.bar(g.Hora-0.2, g.Discharge, width=0.4, color="#E67E22", label="Descarga")
    ax1.bar(g.Hora+0.2, -g.Charge, width=0.4, color="#2980B9", label="Carga")
    ax1.set_xlabel("Hora"); ax1.set_ylabel("Potencia (MW)")
    ax2 = ax1.twinx(); ax2.plot(g.Hora, g.SOC_pct, "-o", ms=3, color="#8E44AD", label="SOC (%)"); ax2.set_ylabel("SOC (%)"); ax2.set_ylim(0,105)
    ax1.set_title("Operación del BESS: carga (–), descarga (+) y SOC"); ax1.set_xticks(range(1,25))
    l1,la1 = ax1.get_legend_handles_labels(); l2,la2 = ax2.get_legend_handles_labels(); ax1.legend(l1+l2, la1+la2, fontsize=8, loc="upper right")
    fn = os.path.join(IMG, f"{case}_bess.png"); plt.savefig(fn); plt.close(); return fn

# ================= PORTADA =================
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("RESULTADOS DEL UNIT COMMITMENT — SISTEMA IEEE 24 BARRAS\nTablas y figuras generadas desde los resultados del modelo")
r.bold = True; r.font.size = Pt(15); r.font.color.rgb = AZUL
doc.add_paragraph()

# ================= 1) TRADICIONAL =================
g, fm, ob, dm = load("Basic")
g_basic = g.copy()
H("Unit Commitment tradicional (sistema IEEE 24 barras)", 1)
tabla_unidades(g, False, False); cap("Tabla 17. Resultados de UC del sistema IEEE 24 Barras tradicional.")
img(fig_despacho("Basic", g, "Despacho térmico horario — UC tradicional"))
cap("Figura 13. Despacho térmico horario y demanda del sistema en el escenario tradicional.")
tabla_flujomax(fm); cap("Tabla 18. Flujos máximos de hora — sistema tradicional.")
tabla_obj(ob); cap("Tabla 19. Desglose de función objetivo — sistema tradicional.")
tabla_dim(dm); cap("Tabla 20. Dimensión del modelo y estado de solución — sistema tradicional.")
doc.add_page_break()

# ================= 2) EÓLICO =================
g, fm, ob, dm = load("Eolico")
H("UC considerando variabilidad de generación eólica pronosticada", 1)
tabla_unidades(g, True, False); cap("Tabla 21. Resultados de UC considerando variabilidad de generación eólica pronosticada.")
img(fig_despacho("Eolico", g, "Despacho horario — escenario con generación eólica", use_wind=True))
cap("Figura 14. Despacho horario apilado de generación térmica y eólica utilizada, y demanda del sistema.")
tabla_eolica(g); cap("Tabla 22. Generación eólica disponible, utilizada y vertimiento — con generación eólica pronosticada.")
tabla_cambios_compromiso(g_basic, g, "eólico")
cap("Tabla 23. Cambios en el compromiso de las unidades térmicas en el escenario con generación eólica respecto al escenario tradicional.")
img(fig_eolica("Eolico", g), 6.0); cap("Figura 15. Energía eólica disponible frente a la utilizada — con generación eólica pronosticada.")
tabla_flujomax(fm); cap("Tabla 23-A. Flujos máximos de hora — considerando generación eólica pronosticada.")
tabla_obj(ob); cap("Tabla 24. Desglose de función objetivo — considerando generación eólica pronosticada.")
tabla_dim(dm); cap("Tabla 25. Dimensión del modelo y estado de solución — considerando generación eólica pronosticada.")
doc.add_page_break()

# ================= 3) BESS =================
g, fm, ob, dm = load("BESS")
H("UC considerando recursos de flexibilidad (BESS)", 1)
tabla_unidades(g, False, True); cap("Tabla 26. Resultados de UC del sistema IEEE 24 Barras considerando recursos flexibles.")
img(fig_despacho("BESS", g, "Despacho horario — escenario con BESS", use_bess=True))
cap("Figura 16. Despacho horario apilado con descarga del BESS, carga negativa y demanda del sistema.")
tabla_bess(g); cap("Tabla 27. Operación del BESS: carga, descarga y SOC — considerando recursos flexibles.")
img(fig_bess("BESS", g), 6.0); cap("Figura 17. Potencia de carga (–) y descarga (+) del BESS y evolución del SOC — con almacenamiento.")
tabla_flujomax(fm); cap("Tabla 28. Flujos máximos de hora — considerando recursos flexibles.")
tabla_cambios_compromiso(g_basic, g, "con BESS")
cap("Tabla 29. Cambios en el compromiso de las unidades térmicas en el escenario con BESS respecto al escenario tradicional.")
tabla_obj(ob); cap("Tabla 29-A. Desglose de función objetivo — considerando recursos flexibles.")
tabla_dim(dm); cap("Tabla 30. Dimensión del modelo y estado de solución — considerando recursos flexibles.")
doc.add_page_break()

# ================= 4) COMPLETO =================
g, fm, ob, dm = load("Final")
H("UC considerando generación eólica pronosticada y recursos de flexibilidad", 1)
tabla_unidades(g, True, True); cap("Tabla 31. Resultados de UC considerando variabilidad de generación eólica pronosticada y recursos flexibles.")
img(fig_despacho("Final", g, "Despacho horario — escenario eólico + BESS", use_wind=True, use_bess=True))
cap("Figura 18. Despacho horario apilado de generación térmica, eólica y descarga del BESS; carga del BESS representada en negativo.")
tabla_eolica(g); cap("Tabla 32. Generación eólica disponible, utilizada y vertimiento — con eólica y recursos flexibles.")
img(fig_eolica("Final", g), 6.0); cap("Figura 19. Energía eólica disponible frente a la utilizada — con eólica y recursos flexibles.")
tabla_bess(g); cap("Tabla 33. Operación del BESS: carga, descarga y SOC — con eólica y recursos flexibles.")
img(fig_bess("Final", g), 6.0); cap("Figura 20. Potencia de carga (–) y descarga (+) del BESS y evolución del SOC — con eólica y recursos flexibles.")
tabla_flujomax(fm); cap("Tabla 34. Flujos máximos de hora — con eólica y recursos flexibles.")
tabla_cambios_compromiso(g_basic, g, "conjunto eólico–BESS")
cap("Tabla 35. Cambios en el compromiso de las unidades térmicas en el escenario conjunto eólico–BESS respecto al escenario tradicional.")
tabla_obj(ob); cap("Tabla 35-A. Desglose de función objetivo — con eólica y recursos flexibles.")
tabla_dim(dm); cap("Tabla 36. Dimensión del modelo y estado de solución — con eólica y recursos flexibles.")
doc.add_page_break()

# ================= 5) UBICACIÓN Y COMPARATIVA =================
H("Ubicación óptima y comparación de escenarios", 1)
sweep = pd.read_csv(os.path.join(LOC, "OUT_sweep.csv")); sweep = sweep[np.isfinite(sweep.costo)]
optc = sweep.costo.min(); top = sweep.sort_values("costo").head(10).reset_index(drop=True)
top["pct"] = 100*(top.costo-optc)/optc
add_table(["#","Nodo BESS","Nodo Eólico","Costo total ($)","Sobrecosto vs. óptimo (%)"],
          [[i+1, int(r.bus_bess), int(r.bus_eolico), fmt(r.costo), f"{fmt(r.pct,3)} %"] for i,r in top.iterrows()],
          fontsz=9, hl={0})
cap("Tabla 37. Diez configuraciones de ubicación del BESS y del parque eólico con el menor costo total de operación.")

fig, ax = plt.subplots(figsize=(7.2,3.4))
labels = [f"B{int(r.bus_bess)}/E{int(r.bus_eolico)}" for _,r in top.iterrows()]
ax.bar(labels, top.pct, color=["#1F3B57" if i==0 else "#5B8FB9" for i in range(len(top))])
ax.set_ylabel("Sobrecosto vs. óptimo (%)"); ax.set_xlabel("Configuración (Nodo BESS / Nodo Eólico)")
ax.set_title("Sobrecosto de las 10 mejores ubicaciones"); plt.xticks(rotation=30, ha="right")
f21 = os.path.join(IMG, "fig21_ubicaciones.png"); plt.savefig(f21); plt.close()
img(f21); cap("Figura 21. Mejores configuraciones de posicionamiento del BESS y la generación eólica.")

res = pd.read_csv(os.path.join(D, "resumen_casos.csv"))
orden = {"Basico":0,"BESS":1,"Eolico":2,"Completo":3}
res["o"] = res.Caso.map(orden); res = res.sort_values("o")
nombres = {"Basico":"Tradicional","BESS":"Tradicional + BESS","Eolico":"Tradicional + Eólico","Completo":"Tradicional + BESS + Eólico"}
base = float(res[res.Caso=="Basico"].Costo.iloc[0])
fig, ax = plt.subplots(figsize=(7.0,3.4))
labs = [nombres[c] for c in res.Caso]; ax.bar(labs, res.Costo, color=["#7f8c8d","#5B8FB9","#2ECC71","#1F3B57"])
ax.set_ylabel("Costo total ($)"); ax.set_title("Costo total de operación por escenario")
ax.set_ylim(res.Costo.min()*0.985, res.Costo.max()*1.005)
for i,(c,v) in enumerate(zip(res.Caso, res.Costo)): ax.text(i, v, fmt(v,2), ha="center", va="bottom", fontsize=7.5)
plt.xticks(rotation=15, ha="right")
f22 = os.path.join(IMG, "fig22_comparativa.png"); plt.savefig(f22); plt.close()
img(f22); cap("Figura 22. Comparación entre los escenarios y su costo total en la función objetivo.")

rows = []
for c in ["Basico","BESS","Eolico","Completo"]:
    v = float(res[res.Caso==c].Costo.iloc[0]); ah = 100*(base-v)/base
    rows.append([nombres[c], fmt(v), "----" if c=="Basico" else f"{fmt(ah,2)} %"])
add_table(["Escenario","Costo Total ($)","Ahorro (%)"], rows, fontsz=9.5, hl={3})
cap("Tabla 38. Ahorro porcentual por escenario respecto del caso tradicional.")

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "GRAFICAS_Y_TABLAS_UC.docx")
try: doc.save(out)
except PermissionError:
    import datetime; out = out.replace(".docx", f"_{datetime.datetime.now():%H%M%S}.docx"); doc.save(out)
print("WORD UC GUARDADO:", out)
