# -*- coding: utf-8 -*-
"""
Genera un documento Word con la grafica del proceso de entrenamiento de la LSTM,
la tabla de valores de error y las 6 graficas de prediccion de 24 horas, a partir
de los resultados en LSTM PREDICTION/OUTPUT (METRICS, PREDICTIONS, FIGURES).
Portable: localiza las carpetas desde su propia ubicacion.
"""
import os, re, glob, pandas as pd
import matplotlib; matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))  # .../LSTM PREDICTION
MET  = os.path.join(BASE, "OUTPUT", "METRICS", "metricas_LSTM_resumen.csv")
PRED = os.path.join(BASE, "OUTPUT", "PREDICTIONS")
FIGS = os.path.join(BASE, "OUTPUT", "FIGURES")
IMG  = os.path.join(os.path.dirname(os.path.abspath(__file__)), "figuras"); os.makedirs(IMG, exist_ok=True)
plt.rcParams.update({"font.size": 9, "axes.grid": True, "grid.alpha": 0.3, "figure.dpi": 130, "savefig.bbox": "tight"})
AZUL = RGBColor(0x1F, 0x3B, 0x57)

doc = Document(); doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(11)

def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexc); tcPr.append(sh)
def setwhite(cell, txt, sz=9, bold=True, fill="1F3B57"):
    cell.text = ""; pr = cell.paragraphs[0]; run = pr.add_run(txt); run.bold = bold; run.font.size = Pt(sz)
    run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); pr.alignment = WD_ALIGN_PARAGRAPH.CENTER; shade(cell, fill)
def H(txt, lvl=1):
    h = doc.add_heading(txt, level=lvl)
    for r in h.runs: r.font.color.rgb = AZUL
def cap(txt):
    p = doc.add_paragraph(); r = p.add_run(txt); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x55,0x55,0x55)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
def img(path, w=6.0):
    doc.add_picture(path, width=Inches(w)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def mes(n): return f"Mes {int(n)}"
def parse_ventana(v):   # "m1...m6 → m7" -> ("Mes 1 – Mes 6", "Mes 7")
    izq, der = v.split("→"); nums_i = re.findall(r"\d+", izq); nums_d = re.findall(r"\d+", der)
    return f"{mes(nums_i[0])} – {mes(nums_i[-1])}", mes(nums_d[0])

# ================= PORTADA =================
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("PRONÓSTICO DE VELOCIDAD DE VIENTO — RED LSTM BIDIRECCIONAL\nGráficas y tablas de resultados")
r.bold = True; r.font.size = Pt(15); r.font.color.rgb = AZUL
doc.add_paragraph()

# ================= 1) PROCESO DE ENTRENAMIENTO =================
H("Proceso de entrenamiento de la red LSTM", 1)
loss_png = os.path.join(FIGS, "loss_entrenamiento.png")
if os.path.exists(loss_png):
    img(loss_png, 6.2)
    cap("Figura 10. Proceso de entrenamiento de la red neuronal LSTM Bidireccional (pérdida MSE por época).")
else:
    doc.add_paragraph("(No se encontró loss_entrenamiento.png; ejecute primero el notebook de la LSTM.)")

# ================= 2) TABLA DE ERROR =================
H("Valores de error calculados", 1)
df = pd.read_csv(MET)
maecol = [c for c in df.columns if c.startswith("MAE")][0]
rmscol = [c for c in df.columns if c.startswith("RMSE")][0]
r2col  = [c for c in df.columns if c.upper().startswith("R2") or c.upper().startswith("R^2") or c == "R2"][0]

t = doc.add_table(rows=2, cols=5); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
# fila 0: "Predicción" (0,0)-(0,1) ; MAE/RMSE/R^2 (0,2..4)-(1,2..4)
a = t.cell(0,0).merge(t.cell(0,1)); setwhite(a, "Predicción")
for j, txt in zip([2,3,4], ["MAE (m/s)", "RMSE (m/s)", "R^2"]):
    c = t.cell(0,j).merge(t.cell(1,j)); setwhite(c, txt)
setwhite(t.cell(1,0), "Input"); setwhite(t.cell(1,1), "Mes Predicción")
for i, row in df.iterrows():
    inp, mp = parse_ventana(str(row["Ventana"]))
    cells = t.add_row().cells
    vals = [inp, mp, f"{float(row[maecol]):.4f}", f"{float(row[rmscol]):.4f}", f"{float(row[r2col]):.4f}"]
    for j, v in enumerate(vals):
        cells[j].text = ""; pr = cells[j].paragraphs[0]; run = pr.add_run(v); run.font.size = Pt(9)
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i % 2 == 1: shade(cells[j], "F2F6FA")
cap("Tabla 16. Valores de error calculados (MAE, RMSE y R² por ventana de predicción).")

# ================= 3) FIGURAS DE PREDICCIÓN =================
H("Predicción de las 24 horas del mes siguiente", 1)
files = glob.glob(os.path.join(PRED, "Prediccion_*.csv"))
def orden(f):
    m = re.search(r"m(\d+)\.\.\.", os.path.basename(f)); return int(m.group(1)) if m else 0
files = sorted(files, key=orden)
fignum = 12
for f in files:
    nm = os.path.basename(f)
    mm = re.search(r"m(\d+)\.\.\.m(\d+)am(\d+)", nm)
    if mm: a, b, c = mm.groups(); titulo = f"Mes {a}...Mes {b} → Predicción Mes {c}"
    else: titulo = nm
    d = pd.read_csv(f)
    fig, ax = plt.subplots(figsize=(7.2, 3.0))
    ax.plot(range(1,25), d["Real"], "-o", ms=3, color="#2C3E50", label="Real")
    ax.plot(range(1,25), d["Prediccion"], "-x", ms=4, color="#E67E22", label="Predicción")
    ax.set_xlabel("Hora del día"); ax.set_ylabel("Viento (m/s)"); ax.set_title(titulo)
    ax.set_xticks(range(1,25)); ax.legend(fontsize=8)
    fn = os.path.join(IMG, f"pred_m{a}_m{b}_m{c}.png"); plt.savefig(fn); plt.close()
    img(fn, 6.0)
    cap(f"Figura {fignum}. Predicción de 24 horas — {titulo}.")
    fignum += 1

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "GRAFICAS_Y_TABLAS_LSTM.docx")
try: doc.save(out)
except PermissionError:
    import datetime; out = out.replace(".docx", f"_{datetime.datetime.now():%H%M%S}.docx"); doc.save(out)
print("WORD LSTM GUARDADO:", out)
